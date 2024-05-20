import spacy
from nltk.tokenize import word_tokenize
import string
import docx
import nltk
from nltk.corpus import stopwords
import re
from nltk.tag import pos_tag
from collections import Counter
import iso639
"""
spacy.cli.download("en_core_web_sm")
nltk.download('stopwords')
spacy.cli.download("es_dep_news_trf")
nltk.download('averaged_perceptron_tagger')
nltk.download('wordnet')
nltk.download('punkt')
"""

from langdetect import detect

#Здесь используем стандарт кодировки ISO 639-1 codes (https://en.wikipedia.org/wiki/List_of_ISO_639_language_codes)
def detect_language(text):
    change = { 
        "en" : "english",
        "sp" : "spain",
        "ru" : "russian",
        "ca" : "catalan"
        }
    try:
        language = detect(text)
        res = iso639.to_name(language)
        return res.lower()
    except:
        return "Не удалось определить язык"


"""
nlp = spacy.load("en_core_web_sm")

def read_docx(file_path, language='english'):
    doc = docx.Document(file_path)
    paragraphs = []
    words = []
    sentences = []
    proper_nouns = []
    places = []
    punctuation = set(string.punctuation)
    stop_words = set(stopwords.words(language))

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            paragraphs.append(paragraph.text)
            # Разбить параграфы на предложения и слова для подсчета
            sentences.extend(re.split(r'[.!?]', paragraph.text))
            words.extend(paragraph.text.split())

    # Удалить пустые строки после разделения на предложения
    sentences = list(filter(None, sentences))

    # Определение и подсчет имен собственных и мест с использованием SpaCy
    for sent in sentences:
        doc = nlp(sent)
        for entity in doc.ents:
            if entity.label_ == "PERSON":
                proper_nouns.append(entity.text)
            elif entity.label_ == "GPE":  # Проверка на место
                places.append(entity.text)

    proper_noun_counts = Counter(proper_nouns)
    place_counts = Counter(places)
    unique_word_count = len(set(words))
    return {
        'text': ' '.join(paragraphs),
        'paragraph_count': len(paragraphs),
        'word_count': len(words),
        'sentence_count': len(sentences),
        'unique_word_count': unique_word_count,
        'names': proper_noun_counts,
        'places': place_counts,
        'sentences': sentences
    }
"""

def read_docx(file_path, language):
    if language.lower() == 'russian':
      nlp = spacy.load("ru_core_news_sm")
    elif language.lower() == 'spanish':
      nlp = spacy.load("es_dep_news_trf")
    elif language.lower() == "english":
      nlp = spacy.load("en_core_web_sm")
    doc = docx.Document(file_path)
    paragraphs = []
    words = []
    sentences = []
    proper_nouns = []
    places = []
    punctuation = set(string.punctuation)
    stop_words = set(stopwords.words(language))

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            paragraphs.append(paragraph.text)
            # Разбить параграфы на предложения и слова для подсчета
            sentences.extend(re.split(r'[.!?]', paragraph.text))
            words.extend(paragraph.text.split())

    # Удалить пустые строки после разделения на предложения
    for sent in sentences:
        doc = nlp(sent)
        for entity in doc.ents:
            if entity.label_ == "PERSON":
                proper_nouns.append(entity.text)
            elif entity.label_ == "GPE":  # Проверка на место
                places.append(entity.text)
    sentences = list(filter(None, sentences))
    
    proper_noun_counts = Counter(proper_nouns)
    place_counts = Counter(places)
    unique_word_count = len(set(words))

    return {
        'text': ' '.join(paragraphs),
        'paragraph_count': len(paragraphs),
        'word_count': len(words),
        'sentence_count': len(sentences),
        'unique_word_count': unique_word_count,
        'names': proper_noun_counts,
        'places': place_counts,
        'sentences': sentences
    }


def preprocess_text(text, language):
    if language.lower() == 'russian':
      nlp = spacy.load("ru_core_news_sm")
    elif language.lower() == 'spanish':
      nlp = spacy.load("es_dep_news_trf")
    elif language.lower() == "english":
      nlp = spacy.load("en_core_web_sm")
    # nlp = es_dep_news_trf.load()
    # Токенизация
    tokens = word_tokenize(text, language=language)

    # Удаление стоп-слов
    punctuation = set(string.punctuation)
    stop_words = set(stopwords.words(language))
    tokens = [word for word in tokens if word.lower() not in stop_words and word not in punctuation and (all(ch.isalpha() or ch.isspace() for ch in word) or word.isdigit())]

    # Лемматизация
    # nlp = es_dep_news_trf.load()
    tokens = [token.lemma_ for token in nlp(' '.join(tokens))]

    # Возвращение предобработанного текста в виде строки
    return ' '.join(tokens)



def tokens(text, language):
    punctuation = set(string.punctuation)
    stop_words = set(stopwords.words(language))
    words = word_tokenize(text.lower())  # Привести к нижнему регистру и токенизировать
    filtered_words = [word for word in text if word.lower() not in stop_words and word not in punctuation and (all(ch.isalpha() or ch.isspace() for ch in word) or word.isdigit())]
    return words
